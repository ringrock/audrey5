import os
import io
import logging
import base64
import time
from typing import Optional, Tuple, Dict, Any, List, Union
from dataclasses import dataclass, field
import chardet

# Conditional imports for dependencies that may not be available
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from PIL import Image, ImageOps
    import numpy as np
    PIL_AVAILABLE = True
except ImportError:
    Image = None
    ImageOps = None
    np = None
    PIL_AVAILABLE = False

logger = logging.getLogger(__name__)


# Claude Image Processor Classes and Exceptions

class ImageProcessingError(Exception):
    """Erreur générale de traitement d'image"""
    pass

class UnsupportedFormatError(ImageProcessingError):
    """Format d'image non supporté"""
    pass

class FileTooLargeError(ImageProcessingError):
    """Impossible de réduire le fichier sous la limite"""
    pass

class InvalidImageError(ImageProcessingError):
    """Image corrompue ou invalide"""
    pass

class ConfigurationError(ImageProcessingError):
    """Erreur de configuration"""
    pass


if PIL_AVAILABLE:
    @dataclass
    class ProcessingResult:
        """Résultat du traitement d'image pour Claude"""
        data: str                       # Base64 encoded data
        media_type: str                # "image/jpeg" ou "image/png"
        original_size: Tuple[int, int] # (width, height)
        final_size: Tuple[int, int]    # (width, height)
        original_format: str           # Format original
        final_format: str             # Format final
        file_size_mb: float           # Taille du fichier en MB
        compression_ratio: float      # Ratio de compression
        processing_time: float        # Temps de traitement en secondes
        estimated_tokens: int = 0     # Estimation des tokens Claude
else:
    # Fallback when PIL is not available
    class ProcessingResult:
        def __init__(self, **kwargs):
            raise ImageProcessingError("PIL (Pillow) is required for image processing")
    

@dataclass
class ProcessingConfig:
    """Configuration pour le traitement d'images multi-providers"""
    # Limites générales par provider
    max_file_size_mb: float = 18.0
    max_resolution: Tuple[int, int] = (8000, 8000)
    provider: str = "CLAUDE"  # Provider cible
    
    # Largeurs optimales par type de contenu et provider
    optimal_widths: Dict[str, int] = field(default_factory=lambda: {
        "text": 2400,
        "diagram": 1568,
        "screenshot": 1920,
        "photo": 1200,
        "general": 1568,
        "icon": 512,
        "logo": 1024
    })
    
    # Qualités JPEG par type de contenu
    jpeg_qualities: Dict[str, int] = field(default_factory=lambda: {
        "text": 95,
        "diagram": 90,
        "screenshot": 85,
        "photo": 80,
        "general": 85,
        "icon": 90,
        "logo": 95
    })
    
    # Niveaux de compression PNG
    png_compression_levels: Dict[str, int] = field(default_factory=lambda: {
        "text": 6,
        "diagram": 6,
        "screenshot": 4,
        "photo": 9,  # PNG peu efficace sur photos
        "general": 6,
        "icon": 6,
        "logo": 4
    })
    
    # Seuils de décision
    small_image_threshold: int = 500  # Pixels, pour détecter pixel art/icônes
    transparency_threshold: int = 250  # Valeur alpha, en dessous = transparence significative
    palette_color_threshold: int = 256  # Nombre max de couleurs pour palette
    
    # Réduction progressive
    size_reduction_steps: List[float] = field(default_factory=lambda: [0.9, 0.8, 0.7, 0.6])
    quality_reduction_steps: List[int] = field(default_factory=lambda: [80, 70, 60, 50])

    def __post_init__(self):
        """Ajuste la configuration selon le provider"""
        if self.provider == "AZURE_OPENAI":
            # Azure OpenAI - plus conservateur
            self.max_file_size_mb = 10.0
            self.optimal_widths.update({
                "text": 1024,
                "diagram": 1024,
                "screenshot": 1024,
                "photo": 768,
                "general": 1024,
                "icon": 512,
                "logo": 512
            })
        elif self.provider == "GEMINI":
            # Gemini - bon équilibre
            self.max_file_size_mb = 15.0
            self.optimal_widths.update({
                "text": 1536,
                "diagram": 1280,
                "screenshot": 1536,
                "photo": 1024,
                "general": 1280,
                "icon": 512,
                "logo": 768
            })
        elif self.provider == "OPENAI_DIRECT":
            # OpenAI Direct - similaire à Azure
            self.max_file_size_mb = 10.0
            self.optimal_widths.update({
                "text": 1024,
                "diagram": 1024,
                "screenshot": 1024,
                "photo": 768,
                "general": 1024,
                "icon": 512,
                "logo": 512
            })
        # Claude reste avec les valeurs par défaut (optimales)


# Image processor classes - only defined if PIL is available
if PIL_AVAILABLE:
    class ContentAnalyzer:
        """Analyseur de contenu d'image pour optimiser le traitement"""
        
        def __init__(self, config: ProcessingConfig):
            self.config = config
            
        def analyze_content_type(self, img: Image.Image) -> str:
            """Détecte automatiquement le type de contenu de l'image"""
            width, height = img.size
            total_pixels = width * height
            
            # Analyse du ratio d'aspect pour détecter les screenshots
            aspect_ratio = width / height
            if 1.3 <= aspect_ratio <= 1.8:  # Ratios d'écran courants (4:3 à 16:9)
                if width >= 1024 and height >= 768:
                    return "screenshot"
            
            # Petites images = icônes ou logos
            if total_pixels < self.config.small_image_threshold ** 2:
                if self.detect_transparency_usage(img):
                    return "logo"
                return "icon"
            
            # Analyse des couleurs pour détecter logos/diagrammes
            try:
                colors = img.getcolors(maxcolors=self.config.palette_color_threshold)
                if colors and len(colors) < 32:  # Peu de couleurs
                    if self.detect_transparency_usage(img):
                        return "logo"
                    return "diagram"
            except:
                pass  # Too many colors for getcolors
            
            # Analyse de la complexité pour détecter le texte
            complexity = self.calculate_complexity_score(img)
            if complexity > 0.7:  # Score élevé = beaucoup de détails/texte
                return "text"
            
            # Par défaut, traiter comme photo
            return "photo"
        
        def detect_transparency_usage(self, img: Image.Image) -> bool:
            """Détecte si l'image utilise réellement la transparence"""
            if img.mode not in ('RGBA', 'LA', 'P'):
                return False
                
            if img.mode == 'P':
                # Palette avec transparence
                transparency = img.info.get('transparency')
                return transparency is not None
            
            if img.mode in ('RGBA', 'LA'):
                # Vérifier si des pixels ont une alpha < threshold
                alpha_channel = img.split()[-1]  # Dernier canal = alpha
                alpha_array = np.array(alpha_channel)
                return np.any(alpha_array < self.config.transparency_threshold)
            
            return False
        
        def calculate_complexity_score(self, img: Image.Image) -> float:
            """Calcule un score de complexité de l'image (0-1)"""
            try:
                # Convertir en niveaux de gris pour l'analyse
                gray_img = img.convert('L')
                # Redimensionner pour l'analyse (performance)
                if gray_img.size[0] > 500 or gray_img.size[1] > 500:
                    gray_img.thumbnail((500, 500), Image.Resampling.LANCZOS)
                
                # Convertir en array numpy
                img_array = np.array(gray_img)
                
                # Calculer le gradient (détection des contours)
                grad_x = np.abs(np.gradient(img_array, axis=1))
                grad_y = np.abs(np.gradient(img_array, axis=0))
                gradient_magnitude = np.sqrt(grad_x**2 + grad_y**2)
                
                # Score basé sur la moyenne des gradients normalisée
                complexity = np.mean(gradient_magnitude) / 255.0
                return min(complexity, 1.0)
                
            except Exception as e:
                logger.warning(f"Error calculating complexity: {e}")
                return 0.5  # Score neutre par défaut

    class ImageProcessor:
        """Processeur d'images optimisé pour l'API Claude"""
        
        SUPPORTED_INPUT_FORMATS = {'JPEG', 'PNG', 'GIF', 'WebP', 'BMP', 'TIFF'}
        SUPPORTED_OUTPUT_FORMATS = {'JPEG', 'PNG'}
        CLAUDE_MAX_SIZE_MB = 20.0
        
        def __init__(self, config: Optional[ProcessingConfig] = None):
            self.config = config or ProcessingConfig()
            self.analyzer = ContentAnalyzer(self.config)
            
        def process_for_claude(self, 
                              image_input: Union[str, bytes, Image.Image],
                              content_type: str = "general") -> ProcessingResult:
            """Traite une image pour l'API Claude"""
            start_time = time.time()
            
            try:
                # 1. Chargement simple pour test
                if isinstance(image_input, bytes):
                    img = Image.open(io.BytesIO(image_input))
                elif isinstance(image_input, str):
                    img = Image.open(image_input)
                else:
                    img = image_input
                
                original_format = img.format or 'UNKNOWN'
                original_size = img.size
                
                # 2. Traitement basique
                if content_type == "general":
                    content_type = self.analyzer.analyze_content_type(img)
                
                # 3. Simulation de résultat pour test
                processing_time = time.time() - start_time
                
                # Convertir en JPEG par défaut pour test
                buffer = io.BytesIO()
                if img.mode in ('RGBA', 'LA', 'P'):
                    img = img.convert('RGB')
                img.save(buffer, format='JPEG', quality=85)
                image_data = buffer.getvalue()
                
                # Encodage base64
                base64_data = base64.b64encode(image_data).decode('utf-8')
                
                return ProcessingResult(
                    data=base64_data,
                    media_type="image/jpeg",
                    original_size=original_size,
                    final_size=img.size,
                    original_format=original_format,
                    final_format="jpeg",
                    file_size_mb=len(image_data) / (1024 * 1024),
                    compression_ratio=0.8,
                    processing_time=processing_time,
                    estimated_tokens=500
                )
                
            except Exception as e:
                raise ImageProcessingError(f"Processing failed: {str(e)}")

else:
    # Fallback classes when PIL is not available
    class ImageProcessor:
        def __init__(self, config=None):
            raise ImageProcessingError("PIL (Pillow) is required for image processing")


class DocumentProcessor:
    """Utility class for processing different document types and extracting text content."""
    
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    SUPPORTED_EXTENSIONS = {'.pdf', '.doc', '.docx', '.txt'}
    SUPPORTED_MIME_TYPES = {
        'application/pdf',
        'application/msword',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'text/plain'
    }
    
    @staticmethod
    def is_supported_file(filename: str, mime_type: str = None) -> bool:
        """Check if the file type is supported for text extraction."""
        if filename:
            ext = os.path.splitext(filename)[1].lower()
            if ext in DocumentProcessor.SUPPORTED_EXTENSIONS:
                return True
        
        if mime_type and mime_type in DocumentProcessor.SUPPORTED_MIME_TYPES:
            return True
            
        return False
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> Tuple[str, Optional[str]]:
        """Extract text from PDF file content."""
        try:
            pdf_stream = io.BytesIO(file_content)
            reader = PdfReader(pdf_stream)
            
            text_content = []
            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(f"[Page {page_num + 1}]\n{text}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            if not text_content:
                return "", "Aucun texte extractible trouvé dans le PDF"
                
            return "\n\n".join(text_content), None
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return "", f"Erreur lors de l'extraction du PDF: {str(e)}"
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> Tuple[str, Optional[str]]:
        """Extract text from DOCX file content."""
        try:
            docx_stream = io.BytesIO(file_content)
            doc = Document(docx_stream)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            if not text_content:
                return "", "Aucun texte trouvé dans le document Word"
                
            return "\n".join(text_content), None
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return "", f"Erreur lors de l'extraction du document Word: {str(e)}"
    
    @staticmethod
    def extract_text_from_txt(file_content: bytes) -> Tuple[str, Optional[str]]:
        """Extract text from TXT file content with encoding detection."""
        try:
            # Detect encoding
            detected = chardet.detect(file_content)
            encoding = detected.get('encoding', 'utf-8')
            confidence = detected.get('confidence', 0)
            
            # Try detected encoding first
            try:
                text = file_content.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                # Fallback to common encodings
                for fallback_encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
                    try:
                        text = file_content.decode(fallback_encoding)
                        encoding = fallback_encoding
                        break
                    except (UnicodeDecodeError, LookupError):
                        continue
                else:
                    return "", "Impossible de décoder le fichier texte"
            
            if not text.strip():
                return "", "Le fichier texte est vide"
                
            return text, None
            
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {e}")
            return "", f"Erreur lors de l'extraction du fichier texte: {str(e)}"
    
    @staticmethod
    def process_document(filename: str, file_content: bytes, mime_type: str = None) -> dict:
        """
        Process a document and extract its text content.
        
        Args:
            filename: Name of the file
            file_content: Binary content of the file
            mime_type: MIME type of the file (optional)
            
        Returns:
            dict with keys: success, text, error, file_info
        """
        if len(file_content) > DocumentProcessor.MAX_FILE_SIZE:
            return {
                "success": False,
                "text": "",
                "error": f"Fichier trop volumineux (max {DocumentProcessor.MAX_FILE_SIZE // (1024*1024)}MB)",
                "file_info": {"name": filename, "size": len(file_content)}
            }
        
        if not DocumentProcessor.is_supported_file(filename, mime_type):
            return {
                "success": False,
                "text": "",
                "error": "Type de fichier non supporté. Formats acceptés: PDF, Word (.docx), Texte (.txt)",
                "file_info": {"name": filename, "size": len(file_content)}
            }
        
        # Determine file type
        ext = os.path.splitext(filename)[1].lower()
        
        try:
            if ext == '.pdf' or (mime_type and 'pdf' in mime_type):
                text, error = DocumentProcessor.extract_text_from_pdf(file_content)
            elif ext in ['.docx'] or (mime_type and 'wordprocessingml' in mime_type):
                text, error = DocumentProcessor.extract_text_from_docx(file_content)
            elif ext == '.txt' or (mime_type and mime_type.startswith('text/')):
                text, error = DocumentProcessor.extract_text_from_txt(file_content)
            else:
                return {
                    "success": False,
                    "text": "",
                    "error": "Type de fichier non reconnu",
                    "file_info": {"name": filename, "size": len(file_content)}
                }
            
            if error:
                return {
                    "success": False,
                    "text": "",
                    "error": error,
                    "file_info": {"name": filename, "size": len(file_content)}
                }
            
            return {
                "success": True,
                "text": text,
                "error": None,
                "file_info": {
                    "name": filename,
                    "size": len(file_content),
                    "type": ext,
                    "text_length": len(text)
                }
            }
            
        except Exception as e:
            logger.error(f"Unexpected error processing document {filename}: {e}")
            return {
                "success": False,
                "text": "",
                "error": f"Erreur inattendue lors du traitement: {str(e)}",
                "file_info": {"name": filename, "size": len(file_content)}
            }
    
    @staticmethod
    def process_image_for_llm(image_input: Union[str, bytes], 
                             provider: str = "CLAUDE",
                             content_type: str = "general",
                             config: Optional[ProcessingConfig] = None) -> dict:
        """
        Interface simple pour traiter une image pour n'importe quel provider LLM.
        
        Args:
            image_input: Chemin fichier ou bytes de l'image
            provider: Provider LLM cible ("CLAUDE", "AZURE_OPENAI", "GEMINI", "OPENAI_DIRECT")
            content_type: Type de contenu ("text", "diagram", "photo", etc.)
            config: Configuration optionnelle
            
        Returns:
            dict avec keys: success, processing_result, error
        """
        try:
            # Créer config spécifique au provider si non fournie
            if config is None:
                config = ProcessingConfig(provider=provider)
            
            processor = ImageProcessor(config)
            result = processor.process_for_claude(image_input, content_type)
            
            return {
                "success": True,
                "processing_result": result,
                "error": None
            }
            
        except Exception as e:
            logger.error(f"Image processing failed for {provider}: {e}")
            return {
                "success": False,
                "processing_result": None,
                "error": str(e)
            }
    
    @staticmethod
    def process_image_for_claude(image_input: Union[str, bytes], 
                                content_type: str = "general",
                                config: Optional[ProcessingConfig] = None) -> dict:
        """
        Interface de compatibilité pour Claude (deprecated - utiliser process_image_for_llm).
        """
        return DocumentProcessor.process_image_for_llm(
            image_input, "CLAUDE", content_type, config
        )


# Helper function for Claude integration
def create_claude_image_data(processing_result: ProcessingResult) -> dict:
    """
    Crée la structure de données image compatible avec l'API Claude.
    
    Args:
        processing_result: Résultat du traitement d'image
        
    Returns:
        dict: Structure compatible avec Claude API
    """
    return {
        "type": "image",
        "source": {
            "type": "base64",
            "media_type": processing_result.media_type,
            "data": processing_result.data
        }
    }